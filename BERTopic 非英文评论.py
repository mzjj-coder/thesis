'''
本文件适用于非英文语料主题建模
2025年3月20日14:46:04
'''

import pandas as pd
import nltk
from nltk.corpus import stopwords
from sentence_transformers import SentenceTransformer
from umap import UMAP
from sklearn.decomposition import PCA
from hdbscan import HDBSCAN
from bertopic.representation import KeyBERTInspired
from bertopic.vectorizers import ClassTfidfTransformer
from bertopic.representation import MaximalMarginalRelevance
import numpy as np
import torch
from bertopic import BERTopic
from sklearn.feature_extraction.text import CountVectorizer
from docx import Document
import os
import re
import jieba
import pickle
from tinysegmenter import TinySegmenter


def clean_text(text):
    text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)# 移除URL
    text = re.sub(r'@[\w]*', '', text) # 移除@用户名
    text = re.sub(r'<.*?>', '', text)    # 移除HTML标签
    text = re.sub(r'[^\w\s]', '', text)  # 移除特殊符号（如果想保留部分符号，可修改）
    return text


# 去除表情符号
def clean_emoji(desstr, restr=''):
        # 过滤表情
        try:
            co = re.compile(u'['u'\U0001F300-\U0001F64F' u'\U0001F680-\U0001F6FF'u'\u2600-\u2B55]+')
        except re.error:
            co = re.compile(u'('u'\ud83c[\udf00-\udfff]|'u'\ud83d[\udc00-\ude4f\ude80-\udeff]|'u'[\u2600-\u2B55])+')
        return co.sub(restr, desstr)

def process_text_by_language(df):
    # 创建clean列的副本
    df['fenci'] = df['clean'].copy()

    # 处理中文(简体)和中文(繁体)
    chinese_mask = (df['Language'] == '中文(简体)') | (df['Language'] == '中文(繁体)')

    # 处理日语
    japanese_mask = df['Language'] == '日语'

    # 对中文文本使用jieba分词
    for idx in df[chinese_mask].index:
        words = jieba.cut(df.loc[idx, 'clean'])
        df.loc[idx, 'fenci'] = ' '.join(words)

    # 对日语文本使用TinySegmenter分词
    segmenter = TinySegmenter()
    for idx in df[japanese_mask].index:
        words = segmenter.tokenize(df.loc[idx, 'clean'])
        df.loc[idx, 'fenci'] = ' '.join(words)

    return df

def preprocess_data(input_file): # 预处理数据
    df = pd.read_csv(input_file)
    df = df[df['text'].apply(lambda x: isinstance(x, str))]  # 过滤掉非字符串的值
    df['text_length'] = df['text'].apply(len)
    df_filtered = df[df['text_length'] >= 10] # 过滤掉长度小于10的评论
    df_filtered['year'] = pd.to_datetime(df_filtered['year'], format='%Y')
    df_filtered['clean'] = df_filtered['text'].apply(lambda x: clean_text(x))
    df_filtered['clean'] = df_filtered['clean'].apply(lambda x: clean_emoji(x))
    df_filtered = process_text_by_language(df_filtered)
    df_new = df_filtered[['year', 'text','clean','fenci']].reset_index(drop=True)
    df_new.dropna(subset = ['fenci'], inplace = True)
    print(len(df_new), "条评论已加载")
    return df_new

def load_stopwords_locally(data_dir='nltk_data'):
        """
        从本地加载停用词，无需网络连接
        """
        # 尝试从预先保存的文件加载
        stopwords_file = os.path.join(os.getcwd(), data_dir, 'all_stopwords.pkl')

        if os.path.exists(stopwords_file):
            with open(stopwords_file, 'rb') as f:
                return pickle.load(f)

        # 如果没有预先保存的文件，尝试直接从NLTK语料库加载
        nltk_data_path = os.path.join(os.getcwd(), data_dir)
        nltk.data.path.append(nltk_data_path)

        languages = stopwords.fileids()
        all_stopwords = set()
        for lang in languages:
            all_stopwords.update(stopwords.words(lang))

        return list(all_stopwords)

# 自定义分词器
def custom_tokenizer(text):
    # 使用基本分词
    words = text.split()
    # 过滤掉长度小于3的词
    return [word for word in words if len(word) >= 3]

def rescale(x, inplace=False):
        """ Rescale an embedding so optimization will not have convergence issues.
        """
        if not inplace:
            x = np.array(x, copy=True)

        x /= np.std(x[:, 0]) * 10000

        return x

def train_topic_model(df_new, output_dir):
    local_model_path = "./paraphrase-multilingual-MiniLM-L12-v2"
    embedding_model = SentenceTransformer(local_model_path)
    embeddings = embedding_model.encode(df_new['fenci'], show_progress_bar=True)
    print('嵌入数组：', embeddings.size)

    if embeddings.size == 0:
        raise ValueError("嵌入数组为空，请检查输入数据。")

    # Initialize and rescale PCA embeddings
    pca_embeddings = rescale(PCA(n_components=5).fit_transform(embeddings))


    # Start UMAP from PCA embeddings
    umap_model = UMAP(
        n_neighbors=15,
        n_components=5,
        min_dist=0.0,
        metric="cosine",
        init=pca_embeddings,
        random_state=42
    )

    # 聚类
    hdbscan_model = HDBSCAN(min_cluster_size=20, metric='euclidean', cluster_selection_method='eom', prediction_data=True)

    vectorizer_model = CountVectorizer(
        stop_words= load_stopwords_locally(),
        min_df=2,
        ngram_range=(1, 2),
        tokenizer=custom_tokenizer
    )

    # Fine-tune your topic representations
    keybert_model = KeyBERTInspired()
    mmr_model = MaximalMarginalRelevance(diversity=0.3)

    representation_model = {
        "KeyBERT": keybert_model,
        "MMR": mmr_model,
    }
    ctfidf_model = ClassTfidfTransformer(reduce_frequent_words=True)

    topic_model = BERTopic(
        embedding_model=embedding_model,
        umap_model=umap_model,
        hdbscan_model=hdbscan_model,
        vectorizer_model=vectorizer_model,
        representation_model=representation_model,
        ctfidf_model=ctfidf_model,
        nr_topics='auto',
        language="multilingual",
        top_n_words=15,
        verbose=True
    )

    topics, probs = topic_model.fit_transform(df_new['fenci'], embeddings)

    doc = Document()
    # 输入标题
    doc.add_heading('topic and frequency', 0)

    doc.add_heading('1.c-TF-IDF', 1)
    t = topic_model.get_topic_info()
    # Iterate over the DataFrame to get topic information
    # 输入基于c-TF-IDF算法的关键词：
    for index, row in t.iterrows():
        topic_number = row['Topic']
        frequency = row['Count']
        keywords = [word for word, _ in topic_model.get_topic(topic_number)]
        # Add a paragraph for each topic with its details
        doc.add_paragraph(f"Topic {topic_number}: {keywords}, Frequency: {frequency}")

    doc.add_heading('2.KeyBERT', 1)
    # 输入基于KeyBERT算法的关键词
    for index, row in t.iterrows():
        topic_number = row['Topic']
        frequency = row['Count']
        keywords = row['KeyBERT']
        doc.add_paragraph(f"Topic {topic_number}: {keywords}, Frequency: {frequency}")

    doc.add_heading('3.MMR', 1)
    for index, row in t.iterrows():
        topic_number = row['Topic']
        frequency = row['Count']
        keywords = row['MMR']
        doc.add_paragraph(f"Topic {topic_number}: {keywords}, Frequency: {frequency}")

    doc.save(os.path.join(output_dir, 'topics_and_frequencies.docx'))

    fig = topic_model.visualize_topics()
    fig.write_html(os.path.join(output_dir, 'results_intertopic_distance.html'))

    model_dir = os.path.join(output_dir, 'model')
    topic_model.save(model_dir, serialization="safetensors", save_ctfidf=True, save_embedding_model=embedding_model)


def process_comments(input_file, output_dir):

    # 预处理数据
    df_new = preprocess_data(input_file)

    # 训练主题模型
    train_topic_model(df_new, output_dir)


# 使用示例
if __name__ == "__main__":
    input_file = './China/comments-China-noEnglish.csv'
    output_dir = './输出文件/China/No_English'
    os.makedirs(output_dir, exist_ok=True)
    process_comments(input_file, output_dir)