'''
本文件适用于英文语料主题建模
2025年3月20日14:46:04
'''

import pandas as pd
import nltk
from nltk.corpus import stopwords
from sentence_transformers import SentenceTransformer
from umap import UMAP
from sklearn.decomposition import PCA
from hdbscan import HDBSCAN
from bertopic.representation import KeyBERTInspired
from bertopic.vectorizers import ClassTfidfTransformer
from bertopic.representation import MaximalMarginalRelevance
import numpy as np
import torch
from bertopic import BERTopic
from sklearn.feature_extraction.text import CountVectorizer
from docx import Document
import os
import pickle



def load_stopwords_locally(data_dir='nltk_data'):
    """
    从本地加载停用词，无需网络连接
    """
    # 尝试从预先保存的文件加载
    stopwords_file = os.path.join(os.getcwd(), data_dir, 'all_stopwords.pkl')

    if os.path.exists(stopwords_file):
        with open(stopwords_file, 'rb') as f:
            return pickle.load(f)

    # 如果没有预先保存的文件，尝试直接从NLTK语料库加载
    nltk_data_path = os.path.join(os.getcwd(), data_dir)
    nltk.data.path.append(nltk_data_path)

    languages = stopwords.fileids()
    all_stopwords = set()
    for lang in languages:
        all_stopwords.update(stopwords.words(lang))

    return list(all_stopwords)

# 自定义分词器
def custom_tokenizer(text):
    # 使用基本分词
    words = text.split()
    # 过滤掉长度小于3的词
    return [word for word in words if len(word) > 3]

def rescale(x, inplace=False):
        """ Rescale an embedding so optimization will not have convergence issues.
        """
        if not inplace:
            x = np.array(x, copy=True)

        x /= np.std(x[:, 0]) * 10000

        return x


def train_topic_model(df_new, output_dir,min_cluster_size=10, use_saved_embeddings=0, topic_n_words=15,nr_topics ='auto'):
    ''''
    use_saved_embeddings是否使用保存的嵌入数组
    '''
    local_model_path = "./paraphrase-multilingual-MiniLM-L12-v2"
    embedding_model = SentenceTransformer(local_model_path)

    if use_saved_embeddings == 0:
        print('开始训练嵌入数组')

        embeddings = embedding_model.encode(df_new['bertopic_clean'], show_progress_bar=True)
        if embeddings.size == 0:
            raise ValueError("嵌入数组为空，请检查输入数据。")
        np.save(os.path.join(output_dir, "document_embeddings.npy"), embeddings)
    else:
        print('使用保存的嵌入数组')
        embeddings = np.load(os.path.join(output_dir, "document_embeddings.npy"))
        if embeddings.size == 0:
            raise ValueError("嵌入数组为空，请检查输入数据。")

    print('嵌入数组：', embeddings.size)

    # Initialize and rescale PCA embeddings
    pca_embeddings = rescale(PCA(n_components=5).fit_transform(embeddings))

    # Start UMAP from PCA embeddings
    umap_model = UMAP(
        n_neighbors=15,
        n_components=5,
        min_dist=0.0,
        metric="cosine",
        init=pca_embeddings,
        random_state=42
    )

    # 聚类
    hdbscan_model = HDBSCAN(min_cluster_size=min_cluster_size, metric='euclidean', cluster_selection_method='eom',
                            prediction_data=True)
    stop_words = load_stopwords_locally()
    stop_words.append('PRON')
    stop_words.append('pron')
    vectorizer_model = CountVectorizer(
        stop_words=stop_words,
        min_df=2,
        ngram_range=(1, 2),
        tokenizer=custom_tokenizer
    )

    # Fine-tune your topic representations
    keybert_model = KeyBERTInspired(top_n_words = topic_n_words)
    mmr_model = MaximalMarginalRelevance(top_n_words = topic_n_words, diversity=0.3)

    representation_model = {
        "KeyBERT": keybert_model,
        "MMR": mmr_model,
    }
    ctfidf_model = ClassTfidfTransformer(reduce_frequent_words=True)

    topic_model = BERTopic(
        embedding_model=embedding_model,
        umap_model=umap_model,
        hdbscan_model=hdbscan_model,
        vectorizer_model=vectorizer_model,
        representation_model=representation_model,
        ctfidf_model=ctfidf_model,
        nr_topics=nr_topics,
        top_n_words=topic_n_words,
        verbose=True,
        language="english"
    )

    topics, probs = topic_model.fit_transform(df_new['bertopic_clean'], embeddings)

    doc = Document()
    # 输入标题
    doc.add_heading('topic and frequency', 0)

    doc.add_heading('1.c-TF-IDF', 1)
    t = topic_model.get_topic_info()
    # Iterate over the DataFrame to get topic information
    # 输入基于c-TF-IDF算法的关键词：
    for index, row in t.iterrows():
        topic_number = row['Topic']
        frequency = row['Count']
        keywords = [word for word, _ in topic_model.get_topic(topic_number)]
        # Add a paragraph for each topic with its details
        doc.add_paragraph(f"Topic {topic_number}: {keywords}, Frequency: {frequency}")

    doc.add_heading('2.KeyBERT', 1)
    # 输入基于KeyBERT算法的关键词
    for index, row in t.iterrows():
        topic_number = row['Topic']
        frequency = row['Count']
        keywords = row['KeyBERT']
        doc.add_paragraph(f"Topic {topic_number}: {keywords}, Frequency: {frequency}")

    doc.add_heading('3.MMR', 1)
    for index, row in t.iterrows():
        topic_number = row['Topic']
        frequency = row['Count']
        keywords = row['MMR']
        doc.add_paragraph(f"Topic {topic_number}: {keywords}, Frequency: {frequency}")

    doc.save(os.path.join(output_dir, 'topics_and_frequencies.docx'))

    fig = topic_model.visualize_topics()
    fig.write_html(os.path.join(output_dir, 'results_intertopic_distance.html'))

    df_new['topic'] = topics

    df_new.to_csv(os.path.join(output_dir, 'documents.csv'), index=False, encoding = 'utf-8-sig')

    model_dir = os.path.join(output_dir, 'model')
    topic_model.save(model_dir, serialization="safetensors", save_ctfidf=True, save_embedding_model=embedding_model)


# 使用示例
if __name__ == "__main__":
    os.chdir(r'C:\Users\mzjj\Desktop\thesis')
    input_file = './France/cleaned_comments.csv'
    output_dir = './输出文件/France/English'
    os.makedirs(output_dir, exist_ok=True)
    # 预处理数据
    df = pd.read_csv(input_file)
    df_new = df.dropna(subset=['bertopic_clean']).reset_index(drop=True)
    df_new = df_new[df_new['bertopic_clean'].apply(lambda x: isinstance(x, str))]  # 过滤掉非字符串的值
    # 训练主题模型，use_saved_embeddings=1表示使用保存的嵌入数组
    train_topic_model(df_new, output_dir, 10, 0, 20, 50)